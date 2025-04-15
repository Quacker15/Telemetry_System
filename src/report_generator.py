import os
import subprocess
from datetime import datetime
from docx import Document

def generate_report(meter_id, power):
    # Ensure the reports directory exists
    reports_dir = os.path.join(os.path.dirname(__file__), '..', 'reports')
    os.makedirs(reports_dir, exist_ok=True)

    # Generate a timestamped filename
    timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    docx_filename = f'{meter_id}_report_{timestamp}.docx'
    docx_path = os.path.join(reports_dir, docx_filename)

    # Create a new Word document
    document = Document()
    document.add_heading(f'Report for {meter_id}', 0)
    document.add_paragraph(f'Report generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    document.add_paragraph(f'Current Power Consumption: {power:.2f} Watts')

    # Save the DOCX file
    try:
        document.save(docx_path)
        print(f'DOCX report saved at: {docx_path}')
    except Exception as e:
        print(f'Error saving DOCX file: {e}')
        return

    # Convert DOCX to PDF using LibreOffice
    try:
        subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', reports_dir,
            docx_path
        ], check=True)
        print(f'PDF report saved at: {os.path.splitext(docx_path)[0]}.pdf')
    except subprocess.CalledProcessError as e:
        print(f'Error converting DOCX to PDF: {e}')
